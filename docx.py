import docx
from pathlib import Path

cwd = Path.cwd() / 'docx'

doc = docx.Document(cwd / 'demo.docx')
len(doc.paragraphs)
doc.paragraphs[0].text

doc.paragraphs[1].text
len(doc.paragraphs[1].runs)
doc.paragraphs[1].runs[0].text
doc.paragraphs[1].runs[1].text
doc.paragraphs[1].runs[2].text
doc.paragraphs[1].runs[3].text

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# Multiple paragraphs
doc = docx.Document()
doc.add_paragraph('Hello world!')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('multipleParagraphs.docx')

# Adding headers
doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')

# Pages and line breaks
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')

# Adding pictures
doc.add_picture('zophie.png', width=docx.shared.Inches(1),
height=docx.shared.Cm(4))